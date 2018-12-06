import docx, os
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
os.chdir('.\\automate_online-materials')
# invitations.docx was previously created with the styles we need.
doc = docx.Document('invitations.docx')
guestsFile = open('guests.txt', 'r')
guests = guestsFile.readlines()
guestsFile.close()
for guest in guests:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firstLine = para.add_run('It would be a pleasure to have the company of')
    firstLine.add_break()
    guestLine = para.add_run(guest)
    guestLine.bold = True
    guestLine.add_break()
    thirdLine = para.add_run('at 11010 Memory Lane on the Evening of April 1st at 7 o\'clock')
    thirdLine.add_break(WD_BREAK.PAGE)
doc.save('invitations.docx')
